"""
Generador de documento Word con formato borgoña/vino tinto.
Estructura por grupo (IP + Proyecto):
  - Portada
  - Por cada grupo: encabezado → por cada vuln: tabla info + tabla plan de acción
"""
from datetime import datetime
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import (
    COLOR_HEADER_BG, COLOR_HEADER_FG,
    COLOR_LABEL_BG, COLOR_LABEL_FG,
    COLOR_BORDER, COLOR_TITLE,
)


def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_table_borders(table, hex_color: str):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), hex_color)
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _cell_text(cell, text: str, bold=False, color_hex: str = "000000",
               font_size: int = 10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text) if text else "")
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = _hex_to_rgb(color_hex)


def _add_header_row(table, col1: str, col2: str):
    row = table.rows[0]
    for cell, text in zip(row.cells, [col1, col2]):
        _set_cell_bg(cell, COLOR_HEADER_BG)
        _cell_text(cell, text, bold=True, color_hex=COLOR_HEADER_FG,
                   font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_info_row(table, label: str, value: str):
    row = table.add_row()
    _set_cell_bg(row.cells[0], COLOR_LABEL_BG)
    _cell_text(row.cells[0], label, bold=True, color_hex=COLOR_LABEL_FG, font_size=10)
    _cell_text(row.cells[1], value or "N/A", color_hex="000000", font_size=10)


def _add_portada(doc: Document, fecha: str):
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PLAN DE REMEDIACIÓN DE VULNERABILIDADES")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = _hex_to_rgb(COLOR_TITLE)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Dirección de Ciberseguridad")
    r.font.size = Pt(13)
    r.font.color.rgb = _hex_to_rgb(COLOR_TITLE)

    doc.add_paragraph()

    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = fecha_p.add_run(f"Generado: {fecha}")
    rf.font.size = Pt(11)
    rf.font.color.rgb = _hex_to_rgb("888888")

    doc.add_page_break()


def _add_group_heading(doc: Document, ip: str, proyecto: str):
    p = doc.add_paragraph()
    run = p.add_run(f"IP: {ip}  |  Proyecto: {proyecto}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = _hex_to_rgb(COLOR_TITLE)
    p.paragraph_format.space_after = Pt(6)


def _add_vuln_block(doc: Document, num: int, vuln: dict):
    # Título de la vulnerabilidad
    nombre = vuln.get("nombre_vulnerabilidad") or "Sin nombre"
    titulo = doc.add_paragraph()
    run = titulo.add_run(f"VULNERABILIDAD {num} – {nombre.upper()}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = _hex_to_rgb(COLOR_TITLE)
    titulo.paragraph_format.space_before = Pt(12)
    titulo.paragraph_format.space_after = Pt(4)

    # ── Tabla 1: Información ──────────────────────────────────────────────────
    t1 = doc.add_table(rows=1, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.LEFT
    t1.columns[0].width = Cm(5)
    t1.columns[1].width = Cm(11)
    _set_table_borders(t1, COLOR_BORDER)
    _add_header_row(t1, "Campo", "Detalle")

    fields_info = [
        ("Nombre del Activo",      vuln.get("nombre_activo")),
        ("Dirección IP",           vuln.get("ip")),
        ("Nombre Vulnerabilidad",  nombre),
        ("Descripción técnica",    vuln.get("descripcion_enriquecida") or vuln.get("descripcion")),
        ("Impacto potencial",      vuln.get("impacto")),
        ("CVE",                    vuln.get("cve")),
        ("Plugin ID (Nessus)",     vuln.get("plugin_id")),
        ("Explotación activa",     vuln.get("explotacion_activa")),
        ("Explotado por malware",  vuln.get("explotado_malware")),
        ("Criticidad técnica",     vuln.get("criticidad")),
        ("CVSS Score",             vuln.get("cvss_score")),
        ("Puerto / Protocolo",     f"{vuln.get('puerto','N/A')} / {vuln.get('protocolo','N/A')}"),
        ("Sistema Operativo",      vuln.get("sistema_operativo")),
        ("Solución detallada",     vuln.get("solucion_detallada") or vuln.get("posible_solucion")),
        ("Estado",                 vuln.get("estado")),
    ]
    for label, value in fields_info:
        _add_info_row(t1, label, value)

    doc.add_paragraph()

    # ── Tabla 2: Plan de Acción ───────────────────────────────────────────────
    plan_title = doc.add_paragraph()
    pr = plan_title.add_run("PLAN DE ACCIÓN")
    pr.bold = True
    pr.font.size = Pt(10)
    pr.font.color.rgb = _hex_to_rgb(COLOR_TITLE)
    plan_title.paragraph_format.space_after = Pt(2)

    t2 = doc.add_table(rows=1, cols=2)
    t2.alignment = WD_TABLE_ALIGNMENT.LEFT
    t2.columns[0].width = Cm(5)
    t2.columns[1].width = Cm(11)
    _set_table_borders(t2, COLOR_BORDER)
    _add_header_row(t2, "Actividad", "Detalle")

    tareas = vuln.get("tareas") or []

    # Campos del plan antes de los pasos
    plan_fields_pre = [
        ("Objetivo",              vuln.get("objetivo")),
        ("Tiempo estimado",       vuln.get("tiempo_estimado")),
        ("Fecha de escaneo",      str(vuln.get("fecha_escaneo", "N/A"))),
        ("Estado vulnerabilidad", vuln.get("estado")),
    ]
    for label, value in plan_fields_pre:
        _add_info_row(t2, label, value)

    # Paso a paso: cada tarea en su propia fila
    if tareas:
        for paso in tareas:
            _add_info_row(t2, "Paso", paso or "N/A")
    else:
        solucion_fb = vuln.get("posible_solucion", "N/A")
        _add_info_row(t2, "Solución", solucion_fb)

    # ── Sección de cierre ─────────────────────────────────────────────────────
    _add_info_row(t2, "Responsable",
                  vuln.get("responsable_sugerido") or "Equipo de Infraestructura / TI")
    _add_info_row(t2, "Prioridad",
                  vuln.get("prioridad") or vuln.get("criticidad") or "Alta")
    _add_info_row(t2, "Fecha objetivo de corrección", "Por definir")

    doc.add_paragraph()


def generar_docx_por_grupo(ip: str, proyecto: str, vulns: list, output_path: str):
    """
    Genera un .docx individual para un grupo (ip + proyecto).
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    fecha = datetime.now().strftime("%d/%m/%Y")
    _add_portada(doc, fecha)
    _add_group_heading(doc, ip, proyecto)

    for i, vuln in enumerate(vulns, 1):
        _add_vuln_block(doc, i, vuln)

    doc.save(output_path)
    print(f"    [OK] Documento: {os.path.basename(output_path)}")
